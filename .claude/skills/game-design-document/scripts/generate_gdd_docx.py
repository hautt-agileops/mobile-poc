#!/usr/bin/env python3
"""Render game_config.json into a professional GDD .docx.

Usage:
    python scripts/generate_gdd_docx.py --config game_config.json --output "Title_GDD_v01.docx"

Needs: python-docx==1.1.2
"""
import argparse

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor, Inches

from _config import load_config, meta_line

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
MUTED = RGBColor(0x60, 0x60, 0x60)


def _style_base(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def _cover(doc, meta):
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(meta["title"])
    r.font.size = Pt(40)
    r.font.bold = True
    r.font.color.rgb = ACCENT
    if meta.get("subtitle"):
        s = doc.add_paragraph()
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = s.add_run(meta["subtitle"])
        sr.font.size = Pt(16)
        sr.font.italic = True
        sr.font.color.rgb = MUTED
    if meta.get("tagline"):
        tg = doc.add_paragraph()
        tg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tg.add_run(meta["tagline"]).font.size = Pt(12)
    for _ in range(4):
        doc.add_paragraph()
    ml = doc.add_paragraph()
    ml.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ml.add_run(meta_line(meta)).font.color.rgb = MUTED
    footer_bits = [b for b in (meta.get("studio"), meta.get("author"),
                               meta.get("date")) if b]
    if footer_bits:
        f = doc.add_paragraph()
        f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        f.add_run("  ·  ".join(footer_bits)).font.color.rgb = MUTED
    doc.add_page_break()


def _add_table(doc, tbl):
    headers = tbl.get("headers", [])
    rows = tbl.get("rows", [])
    if tbl.get("caption"):
        cap = doc.add_paragraph()
        cr = cap.add_run(tbl["caption"])
        cr.font.italic = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = MUTED
    ncols = len(headers) or (len(rows[0]) if rows else 1)
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if headers:
        hcells = table.add_row().cells
        for i, h in enumerate(headers):
            hcells[i].paragraphs[0].add_run(str(h)).bold = True
    for row in rows:
        cells = table.add_row().cells
        for i in range(ncols):
            cells[i].text = str(row[i]) if i < len(row) else ""
    doc.add_paragraph()


def _add_body(doc, block):
    if block.get("body"):
        for para in str(block["body"]).split("\n\n"):
            para = para.strip()
            if para:
                doc.add_paragraph(para)
    for b in block.get("bullets", []):
        doc.add_paragraph(str(b), style="List Bullet")
    for tbl in block.get("tables", []):
        _add_table(doc, tbl)


def build(cfg, output):
    doc = Document()
    _style_base(doc)
    _cover(doc, cfg["meta"])

    doc.add_heading("Table of Contents", level=1)
    for sec in cfg["sections"]:
        doc.add_paragraph(sec.get("heading", "Untitled"), style="List Number")
    doc.add_page_break()

    for sec in cfg["sections"]:
        doc.add_heading(sec.get("heading", "Untitled"), level=1)
        _add_body(doc, sec)
        for sub in sec.get("subsections", []):
            doc.add_heading(sub.get("heading", ""), level=2)
            _add_body(doc, sub)

    doc.save(output)
    print(f"wrote {output}")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--config", required=True)
    ap.add_argument("--output", required=True)
    args = ap.parse_args()
    build(load_config(args.config), args.output)


if __name__ == "__main__":
    main()
